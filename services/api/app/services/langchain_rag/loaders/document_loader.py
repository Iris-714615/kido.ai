"""数据源③：文档加载（PDF / Word / Excel / txt）。

支持：
- PDF：pypdf 提取文本
- Word(.docx)：python-docx 提取段落 + 表格
- Excel(.xlsx)：openpyxl 提取表格并扁平化
- txt：直接读取
表格扁平化、图像做 OCR 占位标记。
"""
from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from app.services.langchain_rag.processors.cleaner import DataCleaner


class DocumentLoader:
    """统一文档加载器，按扩展名分派。"""

    SUPPORTED_EXT = {".pdf", ".docx", ".xlsx", ".txt", ".md"}

    def load(self, file_path: str | Path) -> list[Document]:
        """加载单个文件，返回 Document 列表（一个文件可能产出多个块）。"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXT:
            raise ValueError(f"不支持的文件类型: {ext}，支持 {self.SUPPORTED_EXT}")

        if ext == ".pdf":
            docs = self._load_pdf(path)
        elif ext == ".docx":
            docs = self._load_docx(path)
        elif ext == ".xlsx":
            docs = self._load_xlsx(path)
        else:  # txt / md
            docs = self._load_text(path)

        # 统一注入 source 元数据
        for d in docs:
            d.metadata.setdefault("source", f"document{ext}")
            d.metadata.setdefault("source_id", f"{path.name}_{hash(d.page_content) & 0xffffffff}")
        return DataCleaner.clean_documents(docs)

    def load_directory(self, dir_path: str | Path) -> list[Document]:
        """加载目录下所有支持的文档。"""
        root = Path(dir_path)
        if not root.is_dir():
            raise NotADirectoryError(f"不是目录: {root}")
        docs: list[Document] = []
        for p in root.rglob("*"):
            if p.is_file() and p.suffix.lower() in self.SUPPORTED_EXT:
                try:
                    docs.extend(self.load(p))
                except Exception:
                    continue
        return docs

    # ---------- 各格式实现 ----------
    def _load_text(self, path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [Document(page_content=text, metadata={"file_name": path.name})]

    def _load_pdf(self, path: Path) -> list[Document]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        docs: list[Document] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            # 图像占位：统计页面图片数
            images = getattr(page, "images", []) or []
            if images:
                page_text += "\n" + DataCleaner.image_to_placeholder(f"本页含{len(images)}张图")
            if page_text.strip():
                docs.append(Document(
                    page_content=page_text,
                    metadata={"file_name": path.name, "page": i + 1},
                ))
        return docs

    def _load_docx(self, path: Path) -> list[Document]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        parts: list[str] = []

        # 段落
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 表格扁平化
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            parts.append(DataCleaner.flatten_table(rows))

        # 内嵌图像占位
        try:
            image_count = len(doc.inline_shapes)
            if image_count:
                parts.append(DataCleaner.image_to_placeholder(f"文档含{image_count}张图"))
        except Exception:
            pass

        return [Document(page_content="\n".join(parts), metadata={"file_name": path.name})]

    def _load_xlsx(self, path: Path) -> list[Document]:
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True)
        docs: list[Document] = []
        for ws in wb.worksheets:
            rows: list[list[str]] = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            table_text = DataCleaner.flatten_table(rows)
            if table_text.strip():
                docs.append(Document(
                    page_content=table_text,
                    metadata={"file_name": path.name, "sheet": ws.title},
                ))
        return docs
